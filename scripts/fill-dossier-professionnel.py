#!/usr/bin/env python3
"""Remplit integralement le modele Word officiel du Dossier Professionnel Studi."""

from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

try:
    from docx import Document
except ImportError:
    import subprocess
    import sys

    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document

ROOT = Path(__file__).resolve().parent.parent
SOURCE = Path(
    r"C:\Users\ocean\OneDrive\Documents\Formation WEB DEVELOPPEUR FULL STACK"
    r"\Dossier Professionel\modele dossier pro.docx"
)
OUTPUT = ROOT / "dossier-professionnel" / "DP_Duffour_Oceane.docx"

ADRESSE = "175 imp Charles de Montesquieu, Bât C apt 76, 84200 Carpentras"
VILLE = "Carpentras"
DATE = "18/07/2026"

EXEMPLE_FRONT = {
    "titre": "Conception et développement du parcours commande client (wizard et filtres AJAX)",
    "s1": (
        "Dans le cadre de ma formation Studi, je développe le site e-commerce du traiteur fictif "
        "« Vite & Gourmand ». Pour la partie front-end, je réalise des wireframes desktop et mobile "
        "(accueil, catalogue menus, fiche menu), puis j'intègre les pages statiques en HTML5 sémantique "
        "et Bootstrap 5 (header, footer, pages légales, accessibilité).\n\n"
        "Je développe la partie dynamique : un assistant de commande en 6 étapes sur menu.php "
        "(invités, entrées, plats, desserts, boissons, récapitulatif) et des filtres AJAX sur menus.php. "
        "Je veille à la navigation clavier, aux labels de formulaires et aux attributs ARIA. "
        "Je teste le responsive (375 px) et corrige les alertes WAVE.\n\n"
        "Côté sécurité front-end, j'associe un token CSRF à la soumission du wizard et je complète "
        "la validation JavaScript par une validation serveur PHP."
    ),
    "s2": (
        "HTML5, CSS3, Bootstrap 5.3, JavaScript (Fetch API), Font Awesome, VS Code/Cursor, Git/GitHub, "
        "extension WAVE, Chrome DevTools, charte graphique (#c0392b, #c9a227), MDN Web Docs."
    ),
    "s3": (
        "Je travaille seule sur ce projet fil rouge. J'échange avec mon formateur Studi lors des "
        "corrections de livrables. Je consulte la documentation en ligne (MDN, PHP.net, OWASP)."
    ),
    "s4_org": "Studi — Formation Développeur Web et Web Mobile (projet Vite & Gourmand)",
    "s4_lieu": "Projet de formation — développement web traiteur en ligne",
    "s4_periode": "Du : 01/03/2026    au : 18/07/2026",
    "s5": (
        "J'ai choisi cet exemple car il couvre la compétence front-end : maquettes, intégration statique "
        "et JavaScript dynamique. Le site est déployé sur https://vitegourmand.infinityfree.io/."
    ),
}

EXEMPLE_BACK = {
    "titre": "Conception de la base de données et logique métier des commandes traiteur",
    "s1": (
        "Je modélise la base MySQL : users, menus, plats, menu_options, commandes, commande_details, "
        "commande_historique, avis. J'écris les scripts SQL et les exécute sous phpMyAdmin (XAMPP puis InfinityFree).\n\n"
        "Je développe la couche PDO avec requêtes préparées (includes/db.php, menu-helpers.php). "
        "En local, je synchronise les stats vers MongoDB ; en production, j'adapte le dashboard MySQL.\n\n"
        "J'implémente la logique métier : validation délai livraison, calcul livraison, réduction 10 %, "
        "workflow statuts commande, modération avis, rôles admin/employé/client, bcrypt et CSRF."
    ),
    "s2": "PHP 8.3, MySQL 8, PDO, MongoDB (local), phpMyAdmin, XAMPP, Git, FileZilla, InfinityFree, OWASP.",
    "s3": (
        "Je travaille seule. Mon formateur Studi valide la structure BDD. "
        "Je teste avec les comptes demo jose@, julie@ et client@vite-gourmand.fr."
    ),
    "s4_org": "Studi — Formation Développeur Web et Web Mobile (projet Vite & Gourmand)",
    "s4_lieu": "Back-end PHP/MySQL — site traiteur en ligne",
    "s4_periode": "Du : 01/03/2026    au : 18/07/2026",
    "s5": (
        "Cet exemple couvre les compétences back-end : BDD relationnelle, accès données SQL/NoSQL "
        "et composants métier sur un cas concret de commande traiteur."
    ),
}


def set_cell(row, index: int, text: str) -> None:
    if index < len(row.cells):
        row.cells[index].text = text


def fill_example_table(table, data: dict) -> None:
    set_cell(table.rows[1], 2, data["titre"])
    set_cell(table.rows[5], 2, data["s1"])
    set_cell(table.rows[11], 2, data["s2"])
    set_cell(table.rows[17], 2, data["s3"])
    set_cell(table.rows[24], 2, data["s4_org"])
    set_cell(table.rows[25], 2, data["s4_lieu"])
    set_cell(table.rows[26], 2, data["s4_periode"])
    set_cell(table.rows[30], 2, data["s5"])


def fill_document(doc: Document) -> None:
    # Page de garde
    t0 = doc.tables[0]
    set_cell(t0.rows[1], 2, "DUFFOUR")
    set_cell(t0.rows[2], 2, "DUFFOUR")
    set_cell(t0.rows[3], 2, "Océane Virginie Erica")
    set_cell(t0.rows[4], 2, ADRESSE)

    # Modalité formation
    t1 = doc.tables[1]
    set_cell(t1.rows[6], 1, "☑ Parcours de formation")
    set_cell(t1.rows[7], 1, "☐ Validation des Acquis de l'Expérience (VAE)")

    # Sommaire — exemple 1 par AT
    t3 = doc.tables[3]
    set_cell(t3.rows[3], 1, f"► {EXEMPLE_FRONT['titre']}")
    set_cell(t3.rows[8], 1, f"► {EXEMPLE_BACK['titre']}")

    # Fiches exemples
    fill_example_table(doc.tables[4], EXEMPLE_FRONT)
    fill_example_table(doc.tables[5], EXEMPLE_BACK)

    # Formation
    t6 = doc.tables[6]
    if len(t6.rows) > 4:
        set_cell(t6.rows[4], 0, "Formation Développeur Web et Web Mobile")
        set_cell(t6.rows[4], 1, "Studi")
        set_cell(t6.rows[4], 2, "2025 — 2026")

    # Documents illustrant
    t8 = doc.tables[8]
    if len(t8.rows) > 4:
        set_cell(t8.rows[4], 0, "Capture wizard commande — menu.php?id=14")
    if len(t8.rows) > 5:
        set_cell(t8.rows[5], 0, "Capture dashboard admin — vitegourmand.infinityfree.io")

    # Déclaration — paragraphes après table 7
    for p in doc.paragraphs:
        if "[prénom et nom]" in p.text:
            p.text = p.text.replace(
                "[prénom et nom]",
                "Océane Virginie Erica DUFFOUR",
            )
        if p.text.strip().startswith("Fait à") and "le" in p.text:
            p.text = f"Fait à {VILLE}, le {DATE}"


def export_pdf_via_docx2pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        import docx2pdf  # type: ignore

        docx2pdf.convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception:
        return False


def main() -> int:
    if not SOURCE.exists():
        raise FileNotFoundError(f"Modele introuvable : {SOURCE}")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(str(OUTPUT))
    fill_document(doc)
    doc.save(str(OUTPUT))
    print(f"OK Word complete : {OUTPUT}")

    pdf_out = OUTPUT.with_suffix(".pdf")
    if export_pdf_via_docx2pdf(OUTPUT, pdf_out):
        print(f"OK PDF Word     : {pdf_out}")
    else:
        # Fallback: regenerate from markdown via generate-pdfs if needed
        print("INFO: PDF Word non genere (Word/docx2pdf indisponible). Utilisez export manuel ou DOSSIER_PROFESSIONNEL.pdf")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
